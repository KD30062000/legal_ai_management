import boto3
import asyncio
import os
import tempfile
from typing import List, Dict, Any
import PyPDF2
from docx import Document as DocxDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
import json
from datetime import datetime
from app.models.database import SessionLocal, Document as DBDocument, DocumentChunk
from app.services.vector_store import VectorStore
from app.services.llm_service import LLMService

class DocumentProcessor:
    def __init__(self):
        self.s3_client = boto3.client(
            's3',
            aws_access_key_id=os.getenv("AWS_ACCESS_KEY_ID"),
            aws_secret_access_key=os.getenv("AWS_SECRET_ACCESS_KEY"),
            region_name=os.getenv("S3_REGION")
        )
        self.bucket = os.getenv("S3_BUCKET_NAME")
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
            separators=["\n\n", "\n", ". ", "! ", "? ", " ", ""]
        )
        self.vector_store = VectorStore()
        self.llm_service = LLMService()

    async def process_document(self, document_id: int):
        """Main document processing pipeline"""
        db = SessionLocal()
        try:
            # Get document from database
            doc = db.query(DBDocument).filter(DBDocument.id == document_id).first()
            if not doc:
                raise ValueError(f"Document {document_id} not found")

            # Update status
            doc.processed = "processing"
            db.commit()

            
            # Download from S3
            content = await self._download_from_s3(doc.s3_key)
            
            # Extract text based on file type
            text_content = await self._extract_text(content, doc.content_type, doc.filename)
            
            # Create chunks
            chunks = await self._create_chunks(text_content, doc)
            
            # Generate embeddings and store in vector database
            await self._store_chunks_with_embeddings(chunks, doc, db)
            
            # Generate document summary
            summary = await self._generate_summary(text_content)
            
            # Update document
            doc.processed = "completed"
            doc.processed_at = datetime.utcnow()
            doc.summary = summary
            doc.doc_metadata = {
                "total_chunks": len(chunks),
                "total_characters": len(text_content),
                "processing_completed_at": datetime.utcnow().isoformat()
            }
            db.commit()
            
            return {"status": "success", "chunks_created": len(chunks)}
            
        except Exception as e:
            if 'doc' in locals():
                doc.processed = "failed"
                doc.doc_metadata = {"error": str(e), "failed_at": datetime.utcnow().isoformat()}
                db.commit()
            raise e
        finally:
            db.close()

    async def _download_from_s3(self, s3_key: str) -> bytes:
        """Wait for the uploaded object to appear in S3, then download it.

        This prevents race conditions where the background task starts before
        the client has finished the PUT to the presigned URL.
        """
        max_wait_seconds = int(os.getenv("S3_UPLOAD_WAIT_SECONDS", "120"))
        poll_interval_seconds = 1

        # Poll S3 until the object exists (HEAD succeeds) or timeout
        for _ in range(max_wait_seconds):
            try:
                self.s3_client.head_object(Bucket=self.bucket, Key=s3_key)
                break
            except Exception:
                await asyncio.sleep(poll_interval_seconds)
        else:
            raise Exception(
                f"S3 object {s3_key} not found within {max_wait_seconds} seconds"
            )

        # Download after confirmation
        try:
            response = self.s3_client.get_object(Bucket=self.bucket, Key=s3_key)
            return response["Body"].read()
        except Exception as e:
            raise Exception(f"Failed to download from S3: {str(e)}")

    async def _extract_text(self, content: bytes, content_type: str, filename: str) -> str:
        """Extract text from different file types"""
        try:
            if content_type == "application/pdf":
                return await self._extract_pdf_text(content)
            elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                return await self._extract_docx_text(content)
            elif content_type.startswith("text/"):
                return content.decode('utf-8')
            else:
                raise ValueError(f"Unsupported file type: {content_type}")
        except Exception as e:
            raise Exception(f"Failed to extract text: {str(e)}")

    async def _extract_pdf_text(self, content: bytes) -> str:
        """Extract text from PDF"""
        text = ""
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
            tmp_file.write(content)
            tmp_file.flush()
            
            with open(tmp_file.name, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            
            os.unlink(tmp_file.name)
        return text

    async def _extract_docx_text(self, content: bytes) -> str:
        """Extract text from DOCX"""
        text = ""
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            tmp_file.write(content)
            tmp_file.flush()
            
            doc = DocxDocument(tmp_file.name)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            os.unlink(tmp_file.name)
        return text

    async def _create_chunks(self, text: str, doc: DBDocument) -> List[Document]:
        """Split text into chunks"""
        documents = [Document(
            page_content=text,
            metadata={
                "document_id": doc.id,
                "filename": doc.filename,
                "company_id": doc.company_id,
                "content_type": doc.content_type
            }
        )]
        return self.text_splitter.split_documents(documents)

    async def _store_chunks_with_embeddings(self, chunks: List[Document], doc: DBDocument, db):
        """Store chunks in vector database and create records"""
        # Store in vector database
        chunk_ids = await self.vector_store.add_documents(chunks)
        
        # Create database records for chunks
        for i, (chunk, chunk_id) in enumerate(zip(chunks, chunk_ids)):
            db_chunk = DocumentChunk(
                document_id=doc.id,
                content=chunk.page_content,
                chunk_index=i,
                embedding_id=chunk_id,
                doc_metadata=chunk.metadata
            )
            db.add(db_chunk)
        
        db.commit()

    async def _generate_summary(self, text: str) -> str:
        """Generate document summary using LLM"""
        try:
            prompt = f"""
            Please provide a comprehensive summary of this legal document. Focus on:
            1. Document type and purpose
            2. Key parties involved
            3. Main terms and conditions
            4. Important dates and deadlines
            5. Rights and obligations
            6. Critical legal implications
            
            Document content:
            {text[:4000]}... {'(truncated)' if len(text) > 4000 else ''}
            """
            
            return await self.llm_service.generate_response(prompt)
        except Exception as e:
            return f"Summary generation failed: {str(e)}"